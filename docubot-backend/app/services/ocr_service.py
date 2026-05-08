"""
DocuBot — Servicio OCR.
Motor principal: Tesseract + pdfplumber (extracción nativa de PDFs digitales).
Sin dependencias de servicios cloud: compatible con Railway y desarrollo local.
"""
import asyncio
import io
from dataclasses import dataclass
from typing import List

import pytesseract
from PIL import Image

from app.core.demo_mode import IS_DEMO


@dataclass
class ExtractedPage:
    page_number: int
    text: str
    confidence: float
    ocr_engine: str
    layout_metadata: dict


@dataclass
class ExtractionResult:
    pages: List[ExtractedPage]
    total_pages: int
    used_fallback: bool


class OcrService:
    """
    Extracción de texto de documentos sin servicios cloud.

    Estrategia por tipo:
    - PDF con texto nativo → pdfplumber (rápido, sin OCR)
    - PDF escaneado / imágenes → pdf2image + Tesseract
    - DOCX → python-docx (extracción directa)
    """

    async def extract_from_bytes(
        self, file_bytes: bytes, file_type: str
    ) -> ExtractionResult:
        if IS_DEMO:
            return self._demo_result()

        ft = file_type.lower().lstrip(".")
        if ft == "pdf":
            return await self._extract_pdf(file_bytes)
        elif ft == "docx":
            return await self._extract_docx(file_bytes)
        elif ft in ("png", "jpg", "jpeg", "tiff", "bmp"):
            return await self._extract_image(file_bytes)
        else:
            raise ValueError(f"Tipo de archivo no soportado para OCR: {file_type}")

    # ─── PDF ──────────────────────────────────────────────────────────────
    async def _extract_pdf(self, file_bytes: bytes) -> ExtractionResult:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_pdf_sync, file_bytes)

    def _extract_pdf_sync(self, file_bytes: bytes) -> ExtractionResult:
        # Intento 1: extracción nativa con pdfplumber (PDFs digitales)
        try:
            import pdfplumber
            pages: List[ExtractedPage] = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    tables = page.extract_tables() or []
                    pages.append(ExtractedPage(
                        page_number=i,
                        text=text,
                        confidence=0.98 if len(text.strip()) > 20 else 0.0,
                        ocr_engine="pdfplumber",
                        layout_metadata={"tables_detected": len(tables)},
                    ))

            # Si el texto extraído es suficiente, retornar directamente
            total_text = " ".join(p.text for p in pages)
            if len(total_text.strip()) > 50:
                return ExtractionResult(
                    pages=pages, total_pages=len(pages), used_fallback=False
                )
        except Exception:
            pass

        # Intento 2: OCR con Tesseract para PDFs escaneados
        return self._ocr_pdf_tesseract(file_bytes)

    def _ocr_pdf_tesseract(self, file_bytes: bytes) -> ExtractionResult:
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(file_bytes, dpi=200)
            pages = []
            for i, image in enumerate(images, start=1):
                text = pytesseract.image_to_string(image, lang="spa+eng")
                pages.append(ExtractedPage(
                    page_number=i,
                    text=text,
                    confidence=0.70,
                    ocr_engine="tesseract",
                    layout_metadata={},
                ))
            return ExtractionResult(
                pages=pages, total_pages=len(pages), used_fallback=True
            )
        except Exception as e:
            return ExtractionResult(
                pages=[ExtractedPage(
                    page_number=1, text="",
                    confidence=0.0, ocr_engine="failed",
                    layout_metadata={"error": str(e)},
                )],
                total_pages=1, used_fallback=True,
            )

    # ─── DOCX ─────────────────────────────────────────────────────────────
    async def _extract_docx(self, file_bytes: bytes) -> ExtractionResult:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_docx_sync, file_bytes)

    def _extract_docx_sync(self, file_bytes: bytes) -> ExtractionResult:
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ExtractionResult(
                pages=[ExtractedPage(
                    page_number=1,
                    text=full_text,
                    confidence=0.99,
                    ocr_engine="python-docx",
                    layout_metadata={"paragraphs": len(doc.paragraphs)},
                )],
                total_pages=1, used_fallback=False,
            )
        except Exception as e:
            return ExtractionResult(
                pages=[ExtractedPage(
                    page_number=1, text="",
                    confidence=0.0, ocr_engine="failed",
                    layout_metadata={"error": str(e)},
                )],
                total_pages=1, used_fallback=True,
            )

    # ─── Imagen ───────────────────────────────────────────────────────────
    async def _extract_image(self, file_bytes: bytes) -> ExtractionResult:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_image_sync, file_bytes)

    def _extract_image_sync(self, file_bytes: bytes) -> ExtractionResult:
        image = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(image, lang="spa+eng")
        return ExtractionResult(
            pages=[ExtractedPage(
                page_number=1,
                text=text,
                confidence=0.75,
                ocr_engine="tesseract",
                layout_metadata={},
            )],
            total_pages=1, used_fallback=False,
        )

    # ─── Demo ─────────────────────────────────────────────────────────────
    def _demo_result(self) -> ExtractionResult:
        text = (
            "CONTRATO EPC N° 2023-001\n"
            "Proyecto: Planta Concentradora de Cobre\n\n"
            "Cláusula 15.2 — Penalidades por atraso\n"
            "El Contratista pagará al Mandante una multa equivalente al 0,1% del "
            "valor del contrato por cada día de atraso en la entrega del Hito 3, "
            "con un tope máximo del 10% del monto total.\n\n"
            "Cláusula 8.1 — Plazo de ejecución\n"
            "El plazo total de ejecución es de 24 meses contados desde la fecha "
            "de inicio efectivo de las obras, definida en el Acta de Inicio."
        )
        return ExtractionResult(
            pages=[ExtractedPage(
                page_number=1, text=text,
                confidence=0.99, ocr_engine="demo",
                layout_metadata={},
            )],
            total_pages=1, used_fallback=False,
        )


ocr_service = OcrService()
