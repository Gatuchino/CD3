"""
DocuBot — Servicio de parsing de documentos.
Maneja PDF nativo, DOCX y XLSX sin necesidad de OCR.
"""
import io
from dataclasses import dataclass
from typing import List, Optional
import fitz  # PyMuPDF
import docx
import openpyxl


@dataclass
class ParsedPage:
    page_number: int
    text: str
    is_scanned: bool = False
    layout_metadata: dict = None


@dataclass
class ParseResult:
    pages: List[ParsedPage]
    total_pages: int
    requires_ocr: bool
    file_type: str


class ParserService:

    def parse(self, file_bytes: bytes, file_type: str) -> ParseResult:
        """
        Extrae texto de documentos digitales (no escaneados).
        Detecta si requiere OCR.
        """
        if file_type == "pdf":
            return self._parse_pdf(file_bytes)
        elif file_type == "docx":
            return self._parse_docx(file_bytes)
        elif file_type == "xlsx":
            return self._parse_xlsx(file_bytes)
        else:
            # Para imágenes y PDFs escaneados: delegar al OCR service
            return ParseResult(
                pages=[],
                total_pages=0,
                requires_ocr=True,
                file_type=file_type,
            )

    def _parse_pdf(self, file_bytes: bytes) -> ParseResult:
        """Extrae texto de PDFs digitales con PyMuPDF."""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        scanned_pages = 0

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")

            # Detectar si la página está escaneada (poco texto)
            is_scanned = len(text.strip()) < 50
            if is_scanned:
                scanned_pages += 1

            # Extraer tablas si las hay
            tables = []
            try:
                tabs = page.find_tables()
                for tab in tabs:
                    tables.append({
                        "rows": len(tab.rows),
                        "cols": len(tab.header.cells) if tab.header else 0,
                    })
            except Exception:
                pass

            pages.append(ParsedPage(
                page_number=page_num + 1,
                text=text,
                is_scanned=is_scanned,
                layout_metadata={"tables": tables},
            ))

        doc.close()

        # Si más del 50% de las páginas están escaneadas → requiere OCR
        requires_ocr = scanned_pages > len(pages) * 0.5 if pages else False

        return ParseResult(
            pages=pages,
            total_pages=len(pages),
            requires_ocr=requires_ocr,
            file_type="pdf",
        )

    def _parse_docx(self, file_bytes: bytes) -> ParseResult:
        """Extrae texto de documentos DOCX con python-docx."""
        document = docx.Document(io.BytesIO(file_bytes))
        full_text_parts = []

        for para in document.paragraphs:
            if para.text.strip():
                full_text_parts.append(para.text)

        # Extraer tablas
        table_texts = []
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_texts.append(row_text)

        all_text = "\n".join(full_text_parts)
        if table_texts:
            all_text += "\n\n[TABLAS]\n" + "\n".join(table_texts)

        # DOCX se trata como una sola "página" para el pipeline
        pages = [ParsedPage(
            page_number=1,
            text=all_text,
            is_scanned=False,
            layout_metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
            },
        )]

        return ParseResult(
            pages=pages,
            total_pages=1,
            requires_ocr=False,
            file_type="docx",
        )

    def _parse_xlsx(self, file_bytes: bytes) -> ParseResult:
        """Extrae contenido de hojas de cálculo XLSX."""
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        pages = []

        for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_str = " | ".join(
                    str(cell) for cell in row if cell is not None and str(cell).strip()
                )
                if row_str:
                    rows_text.append(row_str)

            sheet_text = f"[HOJA: {sheet_name}]\n" + "\n".join(rows_text)
            pages.append(ParsedPage(
                page_number=sheet_idx,
                text=sheet_text,
                is_scanned=False,
                layout_metadata={
                    "sheet_name": sheet_name,
                    "max_row": ws.max_row,
                    "max_col": ws.max_column,
                },
            ))

        wb.close()
        return ParseResult(
            pages=pages,
            total_pages=len(pages),
            requires_ocr=False,
            file_type="xlsx",
        )


parser_service = ParserService()
